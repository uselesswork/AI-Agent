from io import BytesIO

import pymupdf
import pytest
from docx import Document

from app.services.document_parser import EmptyDocumentTextError, extract_document_text


def test_extract_utf8_txt() -> None:
    assert extract_document_text("Python 开发实习".encode(), ".txt") == "Python 开发实习"


def test_extract_gb18030_txt() -> None:
    assert extract_document_text("后端开发".encode("gb18030"), ".txt") == "后端开发"


def test_extract_pdf() -> None:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Python internship resume")
    content = document.tobytes()
    document.close()

    assert "Python internship resume" in extract_document_text(content, ".pdf")


def test_extract_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("项目经历")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "FastAPI"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_document_text(buffer.getvalue(), ".docx")

    assert "项目经历" in text
    assert "FastAPI" in text


def test_empty_pdf_has_clear_message() -> None:
    document = pymupdf.open()
    document.new_page()
    content = document.tobytes()
    document.close()

    with pytest.raises(EmptyDocumentTextError, match="OCR"):
        extract_document_text(content, ".pdf")
