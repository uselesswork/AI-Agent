from collections.abc import Callable
from io import BytesIO

import pymupdf
from docx import Document


class DocumentParseError(ValueError):
    """文档损坏或内容无法解析。"""


class EmptyDocumentTextError(DocumentParseError):
    """文档中没有可提取的文本。"""


def _extract_pdf(content: bytes) -> str:
    try:
        with pymupdf.open(stream=content, filetype="pdf") as document:
            return "\n".join(page.get_text("text") for page in document)
    except (pymupdf.FileDataError, RuntimeError, ValueError) as exc:
        raise DocumentParseError("PDF 文件损坏或无法解析。") from exc


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join([*paragraphs, *table_cells])
    except (ValueError, KeyError, OSError) as exc:
        raise DocumentParseError("DOCX 文件损坏或无法解析。") from exc


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("TXT 文件编码无法识别，请使用 UTF-8 编码。")


PARSERS: dict[str, Callable[[bytes], str]] = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_txt,
}


def extract_document_text(content: bytes, extension: str) -> str:
    parser = PARSERS.get(extension.lower())
    if parser is None:
        raise DocumentParseError(f"不支持的文件格式：{extension}")

    text = parser(content).strip()
    if not text:
        raise EmptyDocumentTextError(
            "文档中未提取到文字；如果是扫描版 PDF，请先进行 OCR 文字识别。"
        )
    return text
