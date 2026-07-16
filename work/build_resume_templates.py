from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import BaseDocTemplate, Frame, KeepTogether, PageTemplate, Paragraph, Spacer


OUTPUT_DIR = Path(r"D:\Agent\test")
DOCX_PATH = OUTPUT_DIR / "测试简历模板.docx"
PDF_PATH = OUTPUT_DIR / "测试简历模板.pdf"
TXT_PATH = OUTPUT_DIR / "测试简历模板.txt"
FONT_PATH = Path(r"C:\Windows\Fonts\simhei.ttf")

BLUE = "1F4E79"
GRAY = "666666"
LIGHT_BLUE = "D9EAF7"

PROFILE = {
    "name": "张晨",
    "target": "Python 后端 / AI 应用开发实习生",
    "contact": "138-0000-0000 | zhangchen.test@example.com | 上海",
    "links": "GitHub: github.com/zhangchen-test | 作品集: example.com/zhangchen",
    "summary": "计算机科学本科生，熟悉 Python Web 开发、自动化测试与大模型应用开发。能够独立完成需求拆解、接口设计、数据处理和部署，正在寻找后端开发或 AI 应用开发实习机会。",
}

SECTIONS = [
    (
        "教育经历",
        [
            {
                "title": "华东理工大学 | 计算机科学与技术 | 本科",
                "date": "2023.09 - 2027.06",
                "bullets": ["GPA：3.7/4.0；专业前 20%", "主修课程：数据结构、数据库系统、计算机网络、操作系统、软件工程"],
            }
        ],
    ),
    (
        "专业技能",
        [
            {
                "title": "",
                "date": "",
                "bullets": [
                    "Python：FastAPI、SQLAlchemy、Pydantic、Pytest，具备 RESTful API 开发经验",
                    "数据与工程：MySQL、Redis、Git、Docker、Linux，了解 CI/CD 基本流程",
                    "AI 应用：大模型 API、结构化输出、工具调用、LangGraph、RAG 基本流程",
                    "其他：了解 Java、Go 和前端基础，英语 CET-6",
                ],
            }
        ],
    ),
    (
        "项目经历",
        [
            {
                "title": "AI 求职助手 Agent | 核心开发者",
                "date": "2026.07 - 至今",
                "bullets": [
                    "使用 FastAPI 构建简历上传接口，支持 PDF、DOCX、TXT 文档解析及异常校验",
                    "使用 Pydantic 设计结构化候选人画像与岗位匹配结果，限制模型虚构用户经历",
                    "编写 Pytest 单元测试和接口测试，覆盖空文件、超限文件、无文本 PDF 等边界场景",
                ],
            },
            {
                "title": "校园二手交易平台 | 后端开发",
                "date": "2025.10 - 2025.12",
                "bullets": [
                    "使用 FastAPI、MySQL 实现用户、商品、订单和搜索接口，并生成 OpenAPI 文档",
                    "通过 Redis 缓存热门商品，将本地压测平均响应时间从 180 ms 降至 75 ms",
                ],
            },
        ],
    ),
    (
        "实践经历",
        [
            {
                "title": "校级智能系统实验室 | 学生助理",
                "date": "2025.03 - 2025.09",
                "bullets": [
                    "整理 2,000 余条实验数据，编写 Python 脚本完成清洗、去重和格式转换",
                    "协助维护实验室内部接口文档，并为数据处理脚本补充自动化测试",
                ],
            }
        ],
    ),
    (
        "荣誉与证书",
        [
            {
                "title": "",
                "date": "",
                "bullets": ["校级二等奖学金（2025）", "全国大学生计算机设计大赛校赛二等奖（2025）"],
            }
        ],
    ),
]


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=9.5, bold=False, color="000000") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.05

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Microsoft YaHei"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(9.25)
    bullet.paragraph_format.left_indent = Inches(0.22)
    bullet.paragraph_format.first_line_indent = Inches(-0.15)
    bullet.paragraph_format.space_after = Pt(1.5)
    bullet.paragraph_format.line_spacing = 1.02


def add_docx_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Heading 1"]
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=11.5, bold=True, color=BLUE)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_BLUE)
    paragraph._p.get_or_add_pPr().append(shading)


def add_docx_entry(document: Document, title: str, date: str, bullets: list[str]) -> None:
    if title or date:
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table_widths = (8200, 2300)
        for cell, width in zip(table.rows[0].cells, table_widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        left = table.cell(0, 0).paragraphs[0]
        left.paragraph_format.space_after = Pt(0)
        left.paragraph_format.keep_with_next = True
        set_run_font(left.add_run(title), size=9.7, bold=True)

        right = table.cell(0, 1).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraph_format.space_after = Pt(0)
        right.paragraph_format.keep_with_next = True
        set_run_font(right.add_run(date), size=8.8, color=GRAY)

    for text in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.keep_together = True
        set_run_font(paragraph.add_run(text), size=9.25)


def build_docx() -> None:
    document = Document()
    configure_docx_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    set_run_font(title.add_run(PROFILE["name"]), size=22, bold=True, color=BLUE)

    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target.paragraph_format.space_after = Pt(2)
    set_run_font(target.add_run(PROFILE["target"]), size=11, bold=True)

    for key in ("contact", "links"):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        set_run_font(paragraph.add_run(PROFILE[key]), size=8.8, color=GRAY)

    summary = document.add_paragraph()
    summary.paragraph_format.space_before = Pt(4)
    summary.paragraph_format.space_after = Pt(3)
    set_run_font(summary.add_run(PROFILE["summary"]), size=9.2)

    for heading, entries in SECTIONS:
        add_docx_heading(document, heading)
        for entry in entries:
            add_docx_entry(document, entry["title"], entry["date"], entry["bullets"])

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("测试数据 | 仅用于文档解析功能验证"), size=7.5, color="888888")

    document.core_properties.title = "测试简历模板"
    document.core_properties.subject = "AI 求职助手文档解析测试"
    document.core_properties.author = "AI 求职助手项目"
    document.save(DOCX_PATH)


def pdf_page(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("SimHei", 7.5)
    canvas.setFillColor(colors.HexColor("#888888"))
    canvas.drawCentredString(LETTER[0] / 2, 0.34 * inch, "测试数据 | 仅用于文档解析功能验证")
    canvas.restoreState()


def build_pdf() -> None:
    pdfmetrics.registerFont(TTFont("SimHei", str(FONT_PATH)))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontName="SimHei",
        fontSize=9.1,
        leading=11,
        textColor=colors.HexColor("#111111"),
        spaceAfter=2,
    )
    title = ParagraphStyle(
        "ResumeTitle",
        parent=body,
        fontSize=22,
        leading=25,
        alignment=TA_CENTER,
        textColor=colors.HexColor(f"#{BLUE}"),
        spaceAfter=1,
    )
    centered = ParagraphStyle(
        "ResumeCentered",
        parent=body,
        alignment=TA_CENTER,
        fontSize=8.6,
        leading=10,
        textColor=colors.HexColor(f"#{GRAY}"),
        spaceAfter=1,
    )
    target = ParagraphStyle(
        "ResumeTarget",
        parent=centered,
        fontSize=10.5,
        leading=12,
        textColor=colors.black,
    )
    section_style = ParagraphStyle(
        "ResumeSection",
        parent=body,
        fontSize=11.3,
        leading=13,
        textColor=colors.HexColor(f"#{BLUE}"),
        backColor=colors.HexColor(f"#{LIGHT_BLUE}"),
        borderPadding=(2, 3, 2, 3),
        spaceBefore=4,
        spaceAfter=3,
    )
    entry = ParagraphStyle(
        "ResumeEntry",
        parent=body,
        fontSize=9.4,
        leading=11,
        spaceAfter=1,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body,
        fontSize=9,
        leading=10.7,
        leftIndent=12,
        firstLineIndent=-8,
        spaceAfter=1.5,
    )

    document = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title="测试简历模板",
        author="AI 求职助手项目",
    )
    frame = Frame(document.leftMargin, document.bottomMargin, document.width, document.height, id="resume")
    document.addPageTemplates([PageTemplate(id="resume", frames=[frame], onPage=pdf_page)])

    story = [
        Paragraph(PROFILE["name"], title),
        Paragraph(PROFILE["target"], target),
        Paragraph(PROFILE["contact"], centered),
        Paragraph(PROFILE["links"], centered),
        Spacer(1, 3),
        Paragraph(PROFILE["summary"], body),
    ]

    for heading, entries in SECTIONS:
        story.append(Paragraph(heading, section_style))
        for item in entries:
            block = []
            if item["title"] or item["date"]:
                label = f'<b>{item["title"]}</b><font color="#{GRAY}">　{item["date"]}</font>'
                block.append(Paragraph(label, entry))
            block.extend(Paragraph(f"- {text}", bullet_style) for text in item["bullets"])
            story.append(KeepTogether(block))

    document.build(story)


def build_txt() -> None:
    lines = [
        PROFILE["name"],
        PROFILE["target"],
        PROFILE["contact"],
        PROFILE["links"],
        "",
        "个人简介",
        PROFILE["summary"],
    ]
    for heading, entries in SECTIONS:
        lines.extend(["", heading])
        for item in entries:
            if item["title"] or item["date"]:
                lines.append(f'{item["title"]} | {item["date"]}'.strip(" |"))
            lines.extend(f"- {text}" for text in item["bullets"])
    lines.extend(["", "说明：以上姓名、联系方式、链接和经历均为虚构测试数据。"])
    TXT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_txt()
    build_docx()
    build_pdf()
    for path in (TXT_PATH, DOCX_PATH, PDF_PATH):
        print(f"CREATED {path} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
